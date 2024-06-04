from docx.shared import Mm
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_ROW_HEIGHT_RULE

def MergeTable(table, from_index, to):
    # Функция для объединения ячеек в таблице
    # Объединяем ячейки в столбцах 0, 1 и 2 от from_index до from_index+to
    table.cell(from_index, 0).merge(table.cell(from_index+to, 0))
    table.cell(from_index, 1).merge(table.cell(from_index+to, 1))
    table.cell(from_index, 2).merge(table.cell(from_index+to, 2))

def replace_text_in_paragraphs(paragraphs, data, index):
    # Функция для замены текста в абзацах
    for paragraph in paragraphs:
        for run in paragraph.runs:
            for column_name in data.columns:
                # Проверяем каждый абзац на наличие значения из столбца данных
                if f"{column_name}" in run.text:
                    # Если значение найдено, заменяем его
                    run.text = run.text.replace(f"{column_name}", str(data.at[index, column_name]))


def replacetables(tables, data, index):
    # Функция для замены значений в таблицах
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for column_name in data.columns:
                    # Проверяем каждую ячейку на наличие значения из столбца данных
                    if f"{column_name}" in cell.text:
                        if cell.paragraphs and cell.paragraphs[0].runs:
                            # Если значение найдено, заменяем его и применяем форматирование
                            match column_name:
                                case 'FORM':
                                    cell.text = cell.text.replace(f"{column_name}", str(data.at[index, column_name])).lower()
                                    MakeFMT(cell, 14, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_VERTICAL.TOP, 0, False)

def MakeFMT(cell,
            font_size=11,
            paragraph_Align=WD_ALIGN_PARAGRAPH.CENTER,
            vertical_align=WD_ALIGN_VERTICAL.BOTTOM,
            font_space_after=0,
            bold_type=False):
    # Функция для применения форматирования к ячейке
    cell.vertical_alignment = vertical_align
    rc = cell.paragraphs[0].runs[0]
    for paragraph in cell.paragraphs:
        fmt = paragraph.paragraph_format
        fmt.space_after = Mm(font_space_after)
        paragraph.alignment = paragraph_Align
        for run in paragraph.runs:
            run.bold = bold_type
            run.font.size = Pt(font_size)

def Insert_Str_from(Table, index, disc, ze, mark, alignV):
    # Вставляем информацию о дисциплине в указанную ячейку таблицы
    Table.cell(index, 0).text = disc
    cell = Table.cell(index, 0)
    # Применяем форматирование к ячейке
    MakeFMT(cell, 11, WD_ALIGN_PARAGRAPH.LEFT, alignV, 0, False)
    
    # Обрабатываем информацию о зачётных единицах и вставляем её в таблицу
    textForCell = ze
    textForCell = MakeZE(textForCell)  # Преобразуем информацию о зачётных единицах
    Table.cell(index, 1).text = textForCell
    cell = Table.cell(index, 1)
    # Применяем форматирование к ячейке
    MakeFMT(cell, 11, WD_ALIGN_PARAGRAPH.LEFT, alignV, 0, False)

    # Обрабатываем информацию об оценке и вставляем её в таблицу
    textForCell = mark.lower()  # Приводим оценку к нижнему регистру
    Table.cell(index, 2).text = textForCell
    cell = Table.cell(index, 2)
    # Применяем форматирование к ячейке
    MakeFMT(cell, 11, WD_ALIGN_PARAGRAPH.LEFT, alignV, 0, False)

# Функция для добавления "з.е." к строке с зачётными единицами
def MakeZE(old_str):
    startRec = False
    new_str = ""
    if old_str == "":
        return ""
    for word in old_str:
        if word == ")":
            break
        elif startRec:
            new_str = new_str + word
        elif word == "(":
            startRec = True
    new_str = new_str + " з.е."
    return new_str