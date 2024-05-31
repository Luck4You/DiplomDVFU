from tkinter import *
from tkinter.ttk import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Mm
from docx.shared import Pt
import pandas as pd
import os
import re
from docx import Document

def MakeFMT(cell,
            font_size = 11,
            paragrapg_Align = WD_ALIGN_PARAGRAPH.CENTER,
            vertical_align = WD_ALIGN_VERTICAL.BOTTOM,
            font_space_after = 0,
            bond_type = False):
    cell.vertical_alignment = vertical_align
    rc = cell.paragraphs[0].runs[0]
    for paragraph in cell.paragraphs:
        fmt = paragraph.paragraph_format
        fmt.space_after = Mm(font_space_after)
        paragraph.alignment = paragrapg_Align
        for run in paragraph.runs:
            run.bold = bond_type
            run.font.size = Pt(font_size)

def MergeTable (table, from_index, to):
    table.cell(from_index, 0).merge(table.cell(from_index+to, 0))
    table.cell(from_index, 1).merge(table.cell(from_index+to, 1))
    table.cell(from_index, 2).merge(table.cell(from_index+to, 2))


def replace_text_in_paragraphs(paragraphs, data, index):
    for paragraph in paragraphs:
        for run in paragraph.runs:
            for column_name in data.columns:
                if f"{column_name}" in run.text:
                    run.text = run.text.replace(f"{column_name}", str(data.at[index, column_name]))

def replacetables(tables, data, index):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for column_name in data.columns:
                    if f"{column_name}" in cell.text:
                        if cell.paragraphs and cell.paragraphs[0].runs:
                            match column_name:
                                case 'FORM':
                                    cell.text = cell.text.replace(f"{column_name}", str(data.at[index, column_name])).lower()
                                    MakeFMT(cell,14,WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_VERTICAL.TOP,0,False)

def updateProgressBar(parameter=None):
    if progBar['value'] < 100:
        if parameter is not None:
            if progBar['value'] < 100 and parameter + progBar['value'] > 100:
                progBar['value']=100
            else:
                progBar['value'] += parameter
            progBar['value'] += parameter
        else:
            progBar['value'] += 1
        window.update_idletasks()

def MakeDf():
    global df
    df = df.rename(columns= 
    {
        'Школа студента' : ',MAINSCHOOL',
        'ФИО' : 'FIO',
        'Зачетная книга' : 'ZACHBOOK',
        'Наименование дисциплины' : 'DISCNAME',
        'Школа реализующая дисциплину' : 'DISCSCHOOL',
        'Учебная группа' : 'GROUP',
        'Учебный год' : 'STUDYYEAR',
        'Группа периодов контроля' : 'POLUGOD',
        'Система оценивания' : 'TYPEMARK',
        'Вид контроля' : 'TYPECONTROL',
        'Тип ведомости' : 'TYPEVEDOMOSTI',
        'Оценка' : 'OCENKA',
        'Номер ведомости' : 'NUMVEDOMOSTI',
        'Количество часов (Количество ЗЕТ)' : 'DISCHOURS',
        'Индекс' : 'DISCINDEX',
        'Дата проставления' : 'DATAPROST',
        'Кем проставлена' : 'WHOPROST',
        'Семестр' : 'SEMESTR',
        'Учебный план' : 'STUDYPLAN',
        'Рабочий план' : 'JOBPLAN',
        'Состояние' : 'SOSTOYANIE',
        'Выпускной курс' : 'VIPUSK',
        'Преподаватель ведомости' : 'PREPODVESOMOSTI'
    })
    df = df.drop_duplicates(subset=['ZACHBOOK', 
                                    'DISCNAME', 
                                    'STUDYYEAR', 
                                    'POLUGOD', 
                                    'TYPECONTROL', 
                                    'OCENKA'])
    df = df[df['DISCNAME'] != "Элективные курсы по физической культуре и спорту"]
    df['DISCINDEX_STR'] = df['DISCINDEX'].apply(Sort_by_index) 
    df = df.sort_values(by=['ZACHBOOK', 'DISCINDEX_STR', 'STUDYYEAR', 'POLUGOD'])
    df = df.drop(columns=['DISCINDEX_STR'])
    

# def SortIndexDisc():

def Sort_by_index(name):
    name = name.replace(")", "")
    name = name.replace("(", ".")
    array = str(name).split(".")
    SumNum = 0
    for index, element in enumerate(array):
        if element.isdigit():
            match index:
                case 2:
                    SumNum += int(element)*100
                case 3:
                    SumNum += int(element)*10
            SumNum += int(element)
        elif element == "ФТД":
            SumNum = 99999
            continue
        else:
            match element:
                case "Б":
                    SumNum += 0
                case "О":
                    SumNum += 0
                case "В":
                    SumNum += 10000
                case "ДВ":
                    SumNum += 20000
                case "У":
                    SumNum += 30000
                case "П":
                    SumNum += 40000
                case "Б1":
                    SumNum += 50000
                case "Б2":
                    SumNum += 60000
                case "Б3":
                    SumNum += 70000
    return SumNum


def MakeZE(old_str):
    startRec = False
    new_str = ""
    if old_str == "":
        return ""
    for word in old_str:
        if word == ")": break
        elif startRec: new_str = new_str + word
        elif word == "(": startRec = True
    new_str = new_str + " з.е."
    return new_str

       
def Insert_Str_from (Table, index, disc, ze, mark):
    Table.cell(index, 0).text = disc
    cell = Table.cell(index, 0)
    MakeFMT(cell, 11, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_VERTICAL.BOTTOM, 0, False)
        
    textForCell = ze
    textForCell = MakeZE(textForCell)
    Table.cell(index, 1).text = textForCell
    cell = Table.cell(index, 1)
    MakeFMT(cell, 11, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_VERTICAL.BOTTOM, 0, False)

    textForCell = mark
    textForCell = textForCell.lower()
    Table.cell(index, 2).text = textForCell
    cell = Table.cell(index, 2)
    MakeFMT(cell, 11, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_VERTICAL.BOTTOM, 0, False)


def MakeOne(df_student,file_save):
    doc = Document(shablon_path)

    index_Table = 0
    DiscType_now = "1"
    leftTable = NONE
    rightTable = NONE
    mainTable = NONE
    KursJob = pd.DataFrame(columns=['TYPECONTROL',
                                    'DISCNAME',
                                    'DISCHOURS',
                                    'OCENKA'])
    for table in doc.tables:
        if table.cell(0,0).text == "TABL1TABL":
            leftTable = table
        if table.cell(0,0).text == "TABL2TABL":
            rightTable = table
    mainTable = leftTable
    for index, row in df_student.iterrows():
        if row['TYPECONTROL'] == 'Курсовой проект' or row['TYPECONTROL'] == 'Курсовая работа':
            new_row = {'TYPECONTROL':row['TYPECONTROL'],
                       'DISCNAME':row['DISCNAME'],
                       'DISCHOURS':row['DISCHOURS'],
                       'OCENKA':row['OCENKA']}
            KursJob.loc[len(KursJob.index)] = new_row
        else:
            if not DiscType_now == str(row['DISCINDEX'])[1]:
                if index_Table+(len(discname)//54) > len(mainTable.rows) - 1:
                    mainTable = rightTable
                    index_Table = 0
                MergeTable(mainTable,index_Table,1)
                match(str(row['DISCINDEX'])[1]):
                    case "2":
                        DiscType="Практики\nв том числе:"
                        DiscType_now = "2"
                    case "3":
                        DiscType="Государственная итоговая аттестация\nв том числе:"
                        DiscType_now = "3"
                    case "Т":
                        for jindex, row1 in KursJob.iterrows():
                            discname = f"{row1['TYPECONTROL']}, {row1['DISCNAME']}"
                            if len(discname) >= 54 and index_Table+(len(discname)//54) <= len(mainTable.rows) - 1:
                                MergeTable(mainTable,index_Table,(len(discname)//54))
                                index_Table = index_Table + (len(discname)//54)
                            elif index_Table+(len(discname)//54) > len(mainTable.rows) - 1:
                                mainTable = rightTable
                                index_Table = 0
                            Insert_Str_from(mainTable, 
                                            index_Table, 
                                            discname, 
                                            "", 
                                            str(row1['OCENKA']))
                            index_Table = index_Table + 1
                        #проверка
                        DiscType="Факультативные дисциплины\nв том числе:"
                        DiscType_now = "Т"
                mainTable.cell(index_Table, 0).text = DiscType
                cell = mainTable.cell(index_Table, 0)
                MakeFMT(cell, 11, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_VERTICAL.BOTTOM, 0, False)
                if index_Table + 2 > len(mainTable.rows) - 1:
                    mainTable = rightTable
                    index_Table = 0
                else:
                    index_Table = index_Table + 2
            discname = str(row['DISCNAME'])
            if len(discname) >= 54 and index_Table+(len(discname)//54) <= len(mainTable.rows) - 1:
                MergeTable(mainTable,index_Table,(len(discname)//54))
                index_Table = index_Table + (len(discname)//54)
            elif index_Table+(len(discname)//54) > len(mainTable.rows) - 1:
                mainTable = rightTable
                index_Table = 0
            Insert_Str_from(mainTable, 
                            index_Table, 
                            discname, 
                            str(row['DISCHOURS']), 
                            str(row['OCENKA']))
            if index_Table + 1 > len(mainTable.rows) - 1:
                mainTable = rightTable
                index_Table = 0
            else:
                index_Table = index_Table + 1
    

    #зе проверка
    doc.save(f"C:/Users/pavlov.sa/Desktop/Software/DiplomDVFU-1/Test/Unmerged/{file_save}")

def MakeAll():
    global df
    unique_zachbooks = df['ZACHBOOK'].unique()
    zk_dfs = {zachbook: df[df['ZACHBOOK'] == zachbook] for zachbook in unique_zachbooks}


    for zachbook, df_part in zk_dfs.items():
        filename = f"{zachbook}.docx"
        MakeOne(df_part, filename)
    df.to_excel("C:/Users/pavlov.sa/Desktop/Software/DiplomDVFU-1/Test/SortedDataSet.xlsx")

        
shablon_path = "C:/Users/pavlov.sa/Desktop/Software/DiplomDVFU-1/Shablons/Diplom.docx"
data_file_path = "C:/Users/pavlov.sa/Desktop/Software/DiplomDVFU-1/Test/DataSet.xlsx"
df = pd.read_excel(data_file_path)


window = Tk()  
window.title("Генерация дипломов")
window.geometry('800x400')
window.resizable(False,False)
window.option_add("*tearOff", FALSE)

bar_Variable = IntVar(window, 0)
choosenSetting_Variable = StringVar(window, "Один студент")
saveNameDoc_Variable = StringVar(window, "“ФИО (З.К.)”, как имя документа")

shooseSt_But = Button(window, text="Выбрать студетов")
shooseSt_But.place(x = 30, y = 125)
oneSt_RadioBut = Radiobutton(window, text="Один студент", variable = choosenSetting_Variable , value="Один студент")
oneSt_RadioBut.place(x = 30, y = 150)
oneSt_TextBox = Text(window, height=1, width=20)
oneSt_TextBox.place(x = 30, y = 175)
oneSt_TextBox.insert("1.0", "Зачетная книга")
oneSt_TextBox.tag_add("gray", "1.0", "1.end")
oneSt_TextBox.tag_config("gray", foreground="gray")

group_RadioBut = Radiobutton(window, text="Группа", variable = choosenSetting_Variable , value="Группа")
group_RadioBut.place(x = 200, y = 125)
group_TextBox = Text(window, height=1, width=20)
group_TextBox.place(x = 200, y = 150)
group_TextBox.insert("1.0", "Группа")
group_TextBox.tag_add("gray", "1.0", "1.end")
group_TextBox.tag_config("gray", foreground="gray")

dataSet_RadioBut = Radiobutton(window, text="Весь Dataset", variable = choosenSetting_Variable , value="Весь Dataset")
dataSet_RadioBut.place(x = 365, y = 125)

copyInPDF_ChBox = Checkbutton(window, text = 'Создание копий документов в формате “.pdf”')
copyInPDF_ChBox.place(x = 30, y = 50)

makeFullDocs_ChBox = Checkbutton(window, text = 'Создание объединенного документа')
makeFullDocs_ChBox.place(x = 320, y = 50)
makeFullDocs_But = Button(window, text = 'Расположение', command=updateProgressBar)
makeFullDocs_But.place(x = 320, y = 75)

fioAsName_RadioBut = Radiobutton(window, text="“ФИО (З.К.)”, как имя документа", variable = saveNameDoc_Variable, value="“ФИО (З.К.)”, как имя документа")
fioAsName_RadioBut.place(x = 550, y = 50)
customName_RadioBut = Radiobutton(window, text="Ручное", variable = saveNameDoc_Variable, value="Ручное")
customName_RadioBut.place(x = 550, y = 75)
customName_TextBox = Text(window, height=1, width=20)
customName_TextBox.place(x = 550, y = 100)
customName_TextBox.insert("1.0", "Название документа")
customName_TextBox.tag_add("gray", "1.0", "1.end")
customName_TextBox.tag_config("gray", foreground="gray")

progBar = Progressbar(window, orient="horizontal", length=750, value=0, mode="determinate")
progBar.place(x=25, y=350)

Generate_But = Button(window, text="Сформировать")
Generate_But.place(x=550, y=200)

main_menu = Menu()
 
file_menu = Menu()
file_menu.add_command(label="New", command=MakeDf)
file_menu.add_command(label="Save", command=MakeAll)
file_menu.add_command(label="Open")
file_menu.add_separator()
file_menu.add_command(label="Exit")
 
main_menu.add_cascade(label="File", menu=file_menu)
main_menu.add_cascade(label="Edit")
main_menu.add_cascade(label="View")

window.config(menu=main_menu)
# df = pd.read_excel("C:/Users/pavlov.sa/Desktop/dfe.xlsx")
# df_check = pd.read_excel("C:/Users/pavlov.sa/Desktop/Контингент 17.05.xlsx")

# for index, row in df_check.iterrows():
#     df['Зачётные книги студентов (при необходимости)'] = df['Зачётные книги студентов (при необходимости)'].str.replace(str(row['Зачетная книга']),str(row['ФИО']))

# df.to_excel("C:/Users/pavlov.sa/Desktop/Ready.xlsx")

window.mainloop()